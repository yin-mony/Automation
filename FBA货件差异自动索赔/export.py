"""POP 文档导出：按模板填充 detailInfo 生成 PDF，以及多 SKU 模板检查"""
import calendar
import re
import sys
from copy import deepcopy
from datetime import datetime
from pathlib import Path
from shutil import copyfile

from docx import Document
from docx.oxml import OxmlElement
from docx.oxml.ns import qn

class PopExport:
    """FBA 货件 POP 文档导出与多 SKU 模板检查"""

    @classmethod
    def getBaseDir(cls):
        """获取脚本或 exe 所在目录"""
        if getattr(sys, "frozen", False):
            return Path(sys.executable).resolve().parent
        return Path(__file__).resolve().parent

    @classmethod
    def getResourceDir(cls):
        """获取模板等资源所在目录"""
        if getattr(sys, "frozen", False) and hasattr(sys, "_MEIPASS"):
            return Path(sys._MEIPASS)
        return Path(__file__).resolve().parent

    def __init__(self, baseDir=None):
        # 项目根目录（输出、配置、OCR 临时文件）
        self.baseDir = baseDir or self.getBaseDir()
        # 模板目录始终从资源路径读取（打包后在 _MEIPASS，与 baseDir 分离）
        resourceDir = self.getResourceDir()
        # 单 SKU 模板
        self.templatePath = resourceDir / "db53060fa183_发票模板.docx"
        # 多 SKU 模板
        self.multiTemplatePath = resourceDir / "服务商模板.docx"
        # GUI 选择的 POP 模板，未选择时沿用内置服务商模板
        self.customTemplatePath = None
        # POP 输出目录
        self.exportDir = self.baseDir / "file"
        # 汇总表占位标签
        self.summaryLabels = [
            "Packing List Date:",
            "Shipment ID:",
            "Shipment Name:",
            "SKU Total:",
            "Unit Total:",
        ]
        # 汇总表占位示例值
        self.summaryValues = [
            "12/12/2025",
            "FBA1961YHSPH",
            "FBA ASDN (01/12/2026 01:27)-SBD1",
            "1",
            "132",
        ]
        # 签名区姓名表宽度与左侧 X 坐标
        self.nameTableWidth = 5485
        self.nameTableLeftX = 299
        # 多 SKU 商品表预置数据行数（不含表头）
        self.multiItemPresetRows = 6
        # 多 SKU 商品表模板行高，用于超出预置行时同步下移签名表
        self.multiItemRowHeight = 457
    def setParaText(self, paragraph, text):
        """设置段落文本，尽量保留首 run 样式"""
        if paragraph.runs:
            paragraph.runs[0].text = text
            for run in paragraph.runs[1:]:
                run.text = ""
        else:
            paragraph.text = text

    def setCellText(self, cell, text):
        """设置单元格文本（单段落单元格）"""
        if cell.paragraphs:
            self.setParaText(cell.paragraphs[0], text)
        else:
            cell.text = text

    def setCellParaText(self, cell, paraIndex, text):
        """设置单元格内指定段落文本，保留其余段落与样式结构"""
        if paraIndex < len(cell.paragraphs):
            self.setParaText(cell.paragraphs[paraIndex], text)

    def formatDate(self, createTime):
        """YYYY-MM-DD 减一个月后格式化为 MM/DD/YYYY，无效则返回 None"""
        if not createTime:
            return None
        try:
            dt = datetime.strptime(str(createTime).strip()[:10], "%Y-%m-%d")
        except ValueError:
            return None
        month = dt.month - 1
        year = dt.year
        if month == 0:
            month = 12
            year -= 1
        day = min(dt.day, calendar.monthrange(year, month)[1])
        shifted = datetime(year, month, day)
        return shifted.strftime("%m/%d/%Y")

    def cleanName(self, name):
        """清理 Windows 文件名非法字符"""
        text = str(name or "").strip()
        # 将 Windows 文件名非法字符替换为空格
        text = re.sub(r'[\\/:*?"<>|]+', " ", text)
        # 合并多余空格，避免生成难读文件名
        text = re.sub(r"\s+", " ", text).strip()
        return text or "Unknown"

    def fillSummary(self, summaryTable, shipmentId, name, skuTotal, unitTotal, packingListDate=None):
        """汇总表：多 SKU 模板 5 行按行填值；单 SKU 模板 1 行按段落填值"""
        if len(summaryTable.rows) >= 5:
            if packingListDate:
                self.setCellText(summaryTable.rows[0].cells[1], packingListDate)
            self.setCellText(summaryTable.rows[1].cells[1], shipmentId)
            self.setCellText(summaryTable.rows[2].cells[1], str(name or ""))
            self.setCellText(summaryTable.rows[3].cells[1], str(skuTotal))
            self.setCellText(summaryTable.rows[4].cells[1], str(unitTotal))
            return
        rightCell = summaryTable.rows[0].cells[1]
        if len(rightCell.paragraphs) < 5:
            return
        if packingListDate:
            self.setCellParaText(rightCell, 0, packingListDate)
        self.setCellParaText(rightCell, 1, shipmentId)
        self.setCellParaText(rightCell, 2, str(name or ""))
        self.setCellParaText(rightCell, 3, str(skuTotal))
        self.setCellParaText(rightCell, 4, str(unitTotal))

    def getTableText(self, table):
        """合并表格文本，用于识别模板中的业务表格"""
        texts = []
        for row in table.rows:
            for cell in row.cells:
                texts.append(str(cell.text or "").strip())
        return " ".join(texts)

    def isItemTable(self, table):
        """判断表格是否为 Shipment Items 商品表"""
        if not table.rows:
            return False
        headerText = " ".join(str(cell.text or "").strip() for cell in table.rows[0].cells)
        headerText = re.sub(r"\s+", " ", headerText).lower()
        return (
            "fnsku" in headerText
            and "seller sku" in headerText
            and ("shipped quantity" in headerText or "quantity" in headerText)
        )

    def findItemTable(self, doc):
        """按表头定位商品表，兼容 GUI 上传的成品 POP 模板"""
        for table in doc.tables:
            if self.isItemTable(table):
                return table
        return None

    def findSummaryTable(self, doc):
        """定位包含 Packing List Date 等字段的汇总表"""
        for table in doc.tables:
            if self.isItemTable(table):
                continue
            tableText = self.getTableText(table)
            if "Packing List Date" in tableText and "Shipment ID" in tableText:
                return table
        return None

    def fillAfterLabel(self, doc, label, text, clearNext=False):
        """将指定标签后的第一个非空段落替换为当前货件值"""
        if text is None:
            return False
        for index, para in enumerate(doc.paragraphs):
            if label not in str(para.text or ""):
                continue
            targetIndex = index + 1
            while targetIndex < len(doc.paragraphs) and not doc.paragraphs[targetIndex].text.strip():
                targetIndex += 1
            if targetIndex >= len(doc.paragraphs):
                return False
            self.setParaText(doc.paragraphs[targetIndex], str(text))
            if clearNext:
                nextIndex = targetIndex + 1
                while nextIndex < len(doc.paragraphs) and not doc.paragraphs[nextIndex].text.strip():
                    nextIndex += 1
                if nextIndex < len(doc.paragraphs):
                    nextText = str(doc.paragraphs[nextIndex].text or "").strip()
                    stopWords = ("Amazon ID", "Ship To", "Packing List", "Shipment", "Unit Total")
                    if nextText and not any(word in nextText for word in stopWords):
                        self.setParaText(doc.paragraphs[nextIndex], "")
            return True
        return False

    def fillSummaryParagraphs(self, doc, shipmentId, name, skuTotal, unitTotal, packingListDate=None):
        """汇总字段不是表格时，替换标签区后的五个正文值段落"""
        startIndex = None
        for index, para in enumerate(doc.paragraphs):
            paraText = str(para.text or "")
            if "Packing List Date" in paraText and "Shipment ID" in paraText:
                startIndex = index + 1
                break
        if startIndex is None:
            return False

        valueParas = []
        labelWords = ("Packing List Date", "Shipment ID", "Shipment Name", "SKU Total", "Unit Total")
        for para in doc.paragraphs[startIndex:]:
            paraText = str(para.text or "").strip()
            if "Shipment Items" in paraText:
                break
            if not paraText:
                continue
            if any(word in paraText for word in labelWords):
                continue
            valueParas.append(para)
            if len(valueParas) >= 5:
                break
        if len(valueParas) < 5:
            return False

        values = [
            packingListDate or "",
            shipmentId,
            str(name or ""),
            str(skuTotal),
            str(unitTotal),
        ]
        for para, value in zip(valueParas, values):
            self.setParaText(para, value)
        return True

    def fillCommon(self, doc, detailInfo, shipmentId, skuTotal, unitTotal):
        """填充单条/多条模板共用的 Ship From、Ship To、汇总表字段"""
        source = detailInfo.get("source")
        if source:
            self.fillAfterLabel(doc, "Ship From", str(source), clearNext=True)

        fulfillmentCenterId = str(detailInfo.get("fulfillmentCenterId") or "").strip()
        if fulfillmentCenterId:
            centerCode = fulfillmentCenterId[:4]
            self.fillAfterLabel(doc, "Ship To", f"Amazon Fulfillment Center ({centerCode})")

        packingListDate = self.formatDate(detailInfo.get("createTime"))
        summaryTable = self.findSummaryTable(doc)
        if summaryTable:
            self.fillSummary(
                summaryTable,
                shipmentId,
                detailInfo.get("name"),
                skuTotal,
                unitTotal,
                packingListDate,
            )
        else:
            self.fillSummaryParagraphs(
                doc,
                shipmentId,
                detailInfo.get("name"),
                skuTotal,
                unitTotal,
                packingListDate,
            )

    def fillSingle(self, itemTable, items):
        """单 MSKU 使用服务商模板，只保留表头和一行数据"""
        if not items:
            return
        self.keepRows(itemTable, 1)
        row = itemTable.rows[1]
        it = items[0]
        self.setCellText(row.cells[0], str(it.get("fnSku") or ""))
        self.setCellText(row.cells[1], str(it.get("msku") or ""))
        self.setCellText(row.cells[2], str(it.get("quantity") or ""))

    def keepRows(self, itemTable, dataCount):
        """商品表只保留表头和指定数量的数据行"""
        keepCount = dataCount + 1
        # 从表尾删除多余预留行，避免单 SKU 文件留下空白商品行
        while len(itemTable.rows) > keepCount:
            itemTable._tbl.remove(itemTable.rows[-1]._tr)

    def cloneRow(self, itemTable):
        """复制商品表最后一行（空白数据行）作为样式模板并追加"""
        templateRow = itemTable.rows[-1]
        newTr = deepcopy(templateRow._tr)
        itemTable._tbl.append(newTr)
        return itemTable.rows[-1]

    def ensureRows(self, itemTable, needCount):
        """商品表数据行不足时按模板行 deepcopy 增行，保留行样式"""
        dataRowCount = len(itemTable.rows) - 1
        if needCount <= dataRowCount:
            return
        extra = needCount - dataRowCount
        for _ in range(extra):
            newRow = self.cloneRow(itemTable)
            for cell in newRow.cells:
                self.setCellText(cell, "")

    def removeFloat(self, tbl):
        """移除表格页面绝对定位，使其回到正文流式排版"""
        tblPr = tbl.find(qn("w:tblPr"))
        if tblPr is None:
            return
        for tag in ("tblpPr", "tblOverlap"):
            old = tblPr.find(qn(f"w:{tag}"))
            if old is not None:
                tblPr.remove(old)

    def getRowHeight(self, itemTable):
        """读取多 SKU 商品行高度，模板未设置时使用默认高度"""
        if len(itemTable.rows) < 2:
            return self.multiItemRowHeight
        trPr = itemTable.rows[1]._tr.find(qn("w:trPr"))
        trHeight = trPr.find(qn("w:trHeight")) if trPr is not None else None
        if trHeight is None:
            return self.multiItemRowHeight
        try:
            return int(trHeight.get(qn("w:val")) or self.multiItemRowHeight)
        except (TypeError, ValueError):
            return self.multiItemRowHeight

    def shiftFloatY(self, tbl, offsetY):
        """保持浮动表样式并按指定距离纵向移动"""
        if offsetY == 0:
            return
        tblPr = tbl.find(qn("w:tblPr"))
        tblp = tblPr.find(qn("w:tblpPr")) if tblPr is not None else None
        if tblp is None:
            return
        try:
            oldY = int(tblp.get(qn("w:tblpY")) or 0)
        except (TypeError, ValueError):
            return
        # 只调整 Y 坐标，其余浮动定位、边距、重叠规则保持模板原样
        tblp.set(qn("w:tblpY"), str(oldY + offsetY))

    def shiftSignature(self, doc, itemCount):
        """SKU 超过预置行时保持签名表浮动样式并随商品表下移"""
        if itemCount <= self.multiItemPresetRows:
            return
        if len(doc.tables) < 3:
            return
        extraRows = itemCount - self.multiItemPresetRows
        rowHeight = self.getRowHeight(doc.tables[1])
        # 新增几行商品，就按模板商品行高同步下移右侧签名表
        self.shiftFloatY(doc.tables[2]._tbl, extraRows * rowHeight)

    def shiftSingleSignature(self, doc, itemCount):
        """单 SKU 裁剪商品预留行后，同步上移右侧签名表"""
        if itemCount <= 0 or itemCount >= self.multiItemPresetRows:
            return
        if len(doc.tables) < 3:
            return
        deletedRows = self.multiItemPresetRows - itemCount
        rowHeight = self.getRowHeight(doc.tables[1])
        # 删除几行商品预留行，就按模板商品行高同步上移右侧签名表
        self.shiftFloatY(doc.tables[2]._tbl, -deletedRows * rowHeight)

    def shiftAnchorY(self, anchor, offsetEmu):
        """移动浮动图片锚点的纵向偏移"""
        positionV = anchor.find(qn("wp:positionV"))
        posOffset = positionV.find(qn("wp:posOffset")) if positionV is not None else None
        if posOffset is None:
            return
        try:
            oldValue = int(posOffset.text or 0)
        except (TypeError, ValueError):
            return
        posOffset.text = str(oldValue + offsetEmu)

    def shiftShapeY(self, shape, offsetPoint):
        """移动 VML 浮动文本框的 margin-top"""
        style = shape.get("style") or ""
        match = re.search(r"margin-top:([+-]?(?:\d+(?:\.\d*)?|\.\d+))pt", style)
        if not match:
            return
        oldValue = float(match.group(1))
        newValue = oldValue + offsetPoint
        newStyle = (
            style[:match.start(1)]
            + f"{newValue:.2f}".rstrip("0").rstrip(".")
            + style[match.end(1):]
        )
        shape.set("style", newStyle)

    def shiftFloatingSignature(self, doc, offsetTwips):
        """新模板中签名框是浮动对象，多 SKU 增行时同步下移"""
        if offsetTwips <= 0:
            return
        offsetEmu = int(offsetTwips * 635)
        offsetPoint = offsetTwips / 20
        for anchor in doc._element.body.findall(".//" + qn("wp:anchor")):
            extent = anchor.find(qn("wp:extent"))
            if extent is None:
                continue
            try:
                width = int(extent.get("cx") or 0)
                height = int(extent.get("cy") or 0)
            except (TypeError, ValueError):
                continue
            # 只移动签名图片一类的浮动对象，避免移动页眉横线等细长装饰线
            if width > 1000000 and 200000 < height < 2000000:
                self.shiftAnchorY(anchor, offsetEmu)
        for shape in doc._element.body.findall(".//{urn:schemas-microsoft-com:vml}shape"):
            style = shape.get("style") or ""
            # 签名框是绝对定位文本框，按商品表新增高度整体下移
            if "position:absolute" in style and "margin-top:" in style:
                self.shiftShapeY(shape, offsetPoint)

    def fillMulti(self, itemTable, items):
        """多 MSKU 模板：预置行内填入；超出时复制空白行增行且保留行样式"""
        if not items:
            return
        self.ensureRows(itemTable, len(items))
        dataRows = itemTable.rows[1:]
        for idx, it in enumerate(items):
            row = dataRows[idx]
            self.setCellText(row.cells[0], str(it.get("fnSku") or ""))
            self.setCellText(row.cells[1], str(it.get("msku") or ""))
            self.setCellText(row.cells[2], str(it.get("quantity") or ""))

    def fillItems(self, itemTable, items):
        """按商品表表头写入 FNSKU、Seller SKU、Shipped Quantity"""
        if not items:
            return
        self.ensureRows(itemTable, len(items))
        self.keepRows(itemTable, len(items))

        headers = [str(cell.text or "").strip().lower() for cell in itemTable.rows[0].cells]
        fnIndex = None
        skuIndex = None
        qtyIndex = None
        for index, header in enumerate(headers):
            if "fnsku" in header:
                fnIndex = index
            elif "seller sku" in header:
                skuIndex = index
            elif "quantity" in header:
                qtyIndex = index
        if fnIndex is None or skuIndex is None or qtyIndex is None:
            raise ValueError("POP 商品表缺少 FNSKU、Seller SKU 或 Shipped Quantity 列")

        for index, item in enumerate(items):
            row = itemTable.rows[index + 1]
            for cell in row.cells:
                self.setCellText(cell, "")
            self.setCellText(row.cells[fnIndex], str(item.get("fnSku") or ""))
            self.setCellText(row.cells[skuIndex], str(item.get("msku") or ""))
            self.setCellText(row.cells[qtyIndex], str(item.get("quantity") or ""))

    def toPdf(self, docxPath):
        """将 docx 转为 pdf（Windows 需已安装 Microsoft Word）。"""
        docxPath = Path(docxPath)
        pdfPath = docxPath.with_suffix(".pdf")
        try:
            from docx2pdf import convert
            convert(str(docxPath), str(pdfPath))
        except Exception as e:
            # Word 可能在已保存 PDF 后，于退出 COM 时抛远程过程调用失败
            if pdfPath.is_file() and pdfPath.stat().st_size > 0:
                print(f"DOCX 转 PDF 已生成，但 Word 退出异常，继续使用: {pdfPath}", flush=True)
                return pdfPath
            raise RuntimeError(f"DOCX 转 PDF 失败（需安装 Microsoft Word）: {e}") from e
        if not pdfPath.is_file() or pdfPath.stat().st_size <= 0:
            raise RuntimeError(f"PDF 未生成: {pdfPath}")
        return pdfPath

    def buildPdf(self, detailInfo, templatePath, pdfPath, shipmentId, skuTotal, unitTotal):
        """按 PDF 模板页面叠加当前货件数据，直接生成最终 POP PDF"""
        try:
            import fitz
            import pdfplumber
            from reportlab.lib import colors
            from reportlab.lib.styles import ParagraphStyle
            from reportlab.lib.utils import ImageReader
            from reportlab.pdfbase import pdfmetrics
            from reportlab.pdfbase.cidfonts import UnicodeCIDFont
            from reportlab.platypus import Paragraph, Table, TableStyle
            from reportlab.pdfgen import canvas
            from xml.sax.saxutils import escape
            from io import BytesIO
        except Exception as exc:
            raise RuntimeError(f"PDF 模板导出依赖缺失，请先安装 requirements.txt: {exc}") from exc

        try:
            pdfmetrics.registerFont(UnicodeCIDFont("STSong-Light"))
        except Exception:
            pass

        with pdfplumber.open(str(templatePath)) as pdf:
            page = pdf.pages[0]
            pageWidth = float(page.width)
            pageHeight = float(page.height)
            lines = page.lines

            verticalLines = [
                item for item in lines
                if abs(float(item["x0"]) - float(item["x1"])) < 1
                and float(item["bottom"]) - float(item["top"]) > 10
            ]
            tableGuide = [
                item for item in verticalLines
                if 145 <= float(item["x0"]) <= 155 and float(item["top"]) > 300
            ]
            if tableGuide:
                guide = sorted(tableGuide, key=lambda item: float(item["top"]))[0]
                tableTop = float(guide["top"])
                tableBottom = float(guide["bottom"])
            else:
                tableTop = 349
                tableBottom = 606

            tableLeft = 14.5
            tableRight = 574.5
            tableCol1 = 150.0
            tableCol2 = 506.5
            for item in verticalLines:
                xValue = float(item["x0"])
                if 10 <= xValue <= 20 and abs(float(item["top"]) - tableTop) < 2:
                    tableLeft = xValue
                elif 570 <= xValue <= 580 and abs(float(item["top"]) - tableTop) < 2:
                    tableRight = xValue
                elif 145 <= xValue <= 155 and abs(float(item["top"]) - tableTop) < 2:
                    tableCol1 = xValue
                elif 500 <= xValue <= 512 and abs(float(item["top"]) - tableTop) < 2:
                    tableCol2 = xValue

            signatureLines = [
                item for item in verticalLines
                if float(item["top"]) > tableBottom + 2 and float(item["bottom"]) - float(item["top"]) > 20
            ]
            signatureImageBytes = None
            if signatureLines:
                signatureLeft = min(float(item["x0"]) for item in signatureLines)
                signatureRight = max(float(item["x0"]) for item in signatureLines)
                signatureTop = min(float(item["top"]) for item in signatureLines)
                signatureBottom = max(float(item["bottom"]) for item in signatureLines)
                signatureGap = max(14, signatureTop - tableBottom)
                signatureHeight = signatureBottom - signatureTop
            else:
                signatureLeft = 14.5
                signatureRight = 570.5
                signatureTop = tableBottom + 18
                signatureBottom = signatureTop + 82
                signatureGap = 18
                signatureHeight = 82

            try:
                fitzDoc = fitz.open(str(templatePath))
                fitzPage = fitzDoc[0]
                backgroundPix = fitzPage.get_pixmap(matrix=fitz.Matrix(2, 2), alpha=False)
                backgroundImageBytes = backgroundPix.tobytes("png")
                if signatureLines:
                    clip = fitz.Rect(signatureLeft, signatureTop, signatureRight, signatureBottom)
                    pix = fitzPage.get_pixmap(matrix=fitz.Matrix(2, 2), clip=clip, alpha=False)
                    signatureImageBytes = pix.tobytes("png")
                fitzDoc.close()
            except Exception as exc:
                raise RuntimeError(f"PDF 模板读取失败: {templatePath}") from exc

        pdfCanvas = canvas.Canvas(str(pdfPath), pagesize=(pageWidth, pageHeight))
        pdfCanvas.drawImage(
            ImageReader(BytesIO(backgroundImageBytes)),
            0,
            0,
            width=pageWidth,
            height=pageHeight,
            mask=None,
        )
        pdfCanvas.setFillColor(colors.white)
        pdfCanvas.rect(14, pageHeight - 190, 307, 31, fill=1, stroke=0)
        pdfCanvas.rect(327, pageHeight - 188, 245, 31, fill=1, stroke=0)
        pdfCanvas.rect(94, pageHeight - 297, 230, 83, fill=1, stroke=0)

        items = detailInfo.get("items") or []
        itemCount = max(len(items), 1)
        maxTableBottom = pageHeight - 130
        rowHeight = 23
        if tableTop + (itemCount + 1) * rowHeight + signatureGap + signatureHeight > maxTableBottom:
            rowHeight = max(14, int((maxTableBottom - tableTop - signatureGap - signatureHeight) / (itemCount + 1)))
        tableHeight = (itemCount + 1) * rowHeight
        newTableBottom = tableTop + tableHeight
        coverTableBottom = max(tableBottom, newTableBottom)
        pdfCanvas.rect(tableLeft - 1, pageHeight - coverTableBottom - 2, tableRight - tableLeft + 2, coverTableBottom - tableTop + 4, fill=1, stroke=0)
        pdfCanvas.rect(signatureLeft - 1, pageHeight - signatureBottom - 2, signatureRight - signatureLeft + 2, signatureBottom - signatureTop + 4, fill=1, stroke=0)

        valueStyle = ParagraphStyle(
            "valueStyle",
            fontName="STSong-Light",
            fontSize=8.5,
            leading=10,
            textColor=colors.black,
        )
        cellStyle = ParagraphStyle(
            "cellStyle",
            fontName="STSong-Light",
            fontSize=7.6,
            leading=9,
            textColor=colors.black,
        )
        headerStyle = ParagraphStyle(
            "headerStyle",
            fontName="Helvetica-Bold",
            fontSize=8,
            leading=9,
            textColor=colors.white,
        )

        source = str(detailInfo.get("source") or "")
        sourcePara = Paragraph(escape(source), valueStyle)
        sourcePara.wrapOn(pdfCanvas, 300, 30)
        sourcePara.drawOn(pdfCanvas, 16, pageHeight - 184)

        fulfillmentCenterId = str(detailInfo.get("fulfillmentCenterId") or "").strip()
        centerCode = fulfillmentCenterId[:4]
        shipToText = f"Amazon Fulfillment Center ({centerCode})" if centerCode else "Amazon Fulfillment Center"
        shipToPara = Paragraph(escape(shipToText), valueStyle)
        shipToPara.wrapOn(pdfCanvas, 230, 26)
        shipToPara.drawOn(pdfCanvas, 329, pageHeight - 183)

        pdfCanvas.setFillColor(colors.black)
        pdfCanvas.setFont("STSong-Light", 8.5)
        pdfCanvas.drawString(96, pageHeight - 227, self.formatDate(detailInfo.get("createTime")) or "")
        pdfCanvas.drawString(96, pageHeight - 240, shipmentId)
        namePara = Paragraph(escape(str(detailInfo.get("name") or "")), valueStyle)
        namePara.wrapOn(pdfCanvas, 220, 28)
        namePara.drawOn(pdfCanvas, 96, pageHeight - 265)
        pdfCanvas.drawString(96, pageHeight - 275, str(skuTotal))
        pdfCanvas.drawString(96, pageHeight - 289, str(unitTotal))

        tableData = [[
            Paragraph("FNSKU", headerStyle),
            Paragraph("Seller SKU", headerStyle),
            Paragraph("Shipped<br/>Quantity", headerStyle),
        ]]
        if items:
            for item in items:
                tableData.append([
                    Paragraph(escape(str(item.get("fnSku") or "")), cellStyle),
                    Paragraph(escape(str(item.get("msku") or "")), cellStyle),
                    Paragraph(escape(str(item.get("quantity") or "")), cellStyle),
                ])
        else:
            tableData.append([
                Paragraph("", cellStyle),
                Paragraph("", cellStyle),
                Paragraph("", cellStyle),
            ])

        colWidths = [tableCol1 - tableLeft, tableCol2 - tableCol1, tableRight - tableCol2]
        rowHeights = [rowHeight] + [rowHeight for _ in range(itemCount)]
        itemTable = Table(tableData, colWidths=colWidths, rowHeights=rowHeights)
        itemTable.setStyle(TableStyle([
            ("BACKGROUND", (0, 0), (-1, 0), colors.black),
            ("TEXTCOLOR", (0, 0), (-1, 0), colors.white),
            ("FONTNAME", (0, 0), (-1, 0), "Helvetica-Bold"),
            ("FONTNAME", (0, 1), (-1, -1), "STSong-Light"),
            ("GRID", (0, 0), (-1, -1), 0.5, colors.HexColor("#8c8c8c")),
            ("VALIGN", (0, 0), (-1, -1), "MIDDLE"),
            ("ALIGN", (2, 0), (2, -1), "CENTER"),
            ("LEFTPADDING", (0, 0), (-1, -1), 4),
            ("RIGHTPADDING", (0, 0), (-1, -1), 4),
            ("TOPPADDING", (0, 0), (-1, -1), 2),
            ("BOTTOMPADDING", (0, 0), (-1, -1), 2),
        ]))
        itemTable.wrapOn(pdfCanvas, tableRight - tableLeft, tableHeight)
        itemTable.drawOn(pdfCanvas, tableLeft, pageHeight - newTableBottom)

        newSignatureTop = newTableBottom + signatureGap
        newSignatureBottom = newSignatureTop + signatureHeight
        if newSignatureBottom > pageHeight - 40:
            newSignatureBottom = pageHeight - 40
            newSignatureTop = newSignatureBottom - signatureHeight
        if signatureImageBytes:
            pdfCanvas.drawImage(
                ImageReader(BytesIO(signatureImageBytes)),
                signatureLeft,
                pageHeight - newSignatureBottom,
                width=signatureRight - signatureLeft,
                height=signatureHeight,
                mask=None,
            )
        else:
            pdfCanvas.setStrokeColor(colors.black)
            pdfCanvas.rect(signatureLeft, pageHeight - newSignatureBottom, signatureRight - signatureLeft, signatureHeight, fill=0, stroke=1)

        pdfCanvas.save()
        if not pdfPath.is_file() or pdfPath.stat().st_size <= 0:
            raise RuntimeError(f"PDF 未生成: {pdfPath}")
        return pdfPath

    def build(self, detailInfo):
        """按模板填充 detailInfo，生成 POP PDF 文档，调试时可保留 DOCX"""
        items = detailInfo.get("items") or []
        # SKU 总数
        skuTotal = len(items)
        # 装箱总数量
        unitTotal = 0
        for it in items:
            try:
                unitTotal += int(it.get("quantity") or 0)
            except (TypeError, ValueError):
                pass

        shopName = str(detailInfo.get("shopName") or "Unknown").strip()
        shipmentId = str(detailInfo.get("shipmentId") or "").strip()
        # 货件编号为空时无法生成可追踪的 POP 文件
        if not shipmentId:
            raise ValueError("货件编号不能为空，无法生成 POP")

        baseName = self.cleanName(f"{shopName}_{shipmentId}_POP")
        self.exportDir.mkdir(parents=True, exist_ok=True)
        docxPath = self.exportDir / f"{baseName}.docx"
        pdfPath = self.exportDir / f"{baseName}.pdf"

        # 调试查看版式时可只保留 docx，不进入 PDF 转换
        keepDocx = bool(detailInfo.get("keepDocx"))
        templatePath = Path(detailInfo.get("templatePath") or self.customTemplatePath or self.multiTemplatePath)
        # 模板不存在时提前报错，避免 copyfile 抛出不直观异常
        if not templatePath.is_file():
            raise FileNotFoundError(f"POP 模板不存在: {templatePath}")
        if templatePath.suffix.lower() == ".pdf":
            return str(self.buildPdf(detailInfo, templatePath, pdfPath, shipmentId, skuTotal, unitTotal))
        if templatePath.suffix.lower() != ".docx":
            raise ValueError(f"POP 模板仅支持 .docx 或 .pdf: {templatePath}")

        # 复制模板到临时 docx，填充后转 PDF
        copyfile(templatePath, docxPath)
        doc = Document(str(docxPath))
        # POP 模板至少需要汇总表与商品表
        itemTable = self.findItemTable(doc)
        if itemTable is None:
            raise ValueError(f"POP 模板未找到 FNSKU/Seller SKU/Shipped Quantity 商品表: {templatePath}")
        templateDataRows = max(len(itemTable.rows) - 1, 1)
        itemRowHeight = self.getRowHeight(itemTable)

        # 当前货件公共信息写入模板正文或汇总表，商品明细按表头写入商品表
        self.fillCommon(doc, detailInfo, shipmentId, skuTotal, unitTotal)
        self.fillItems(itemTable, items)
        if len(doc.tables) >= 3:
            if len(items) > self.multiItemPresetRows:
                self.shiftSignature(doc, len(items))
            elif len(items) < self.multiItemPresetRows:
                self.shiftSingleSignature(doc, len(items))
        else:
            extraRows = len(items) - templateDataRows
            self.shiftFloatingSignature(doc, extraRows * itemRowHeight)

        doc.save(str(docxPath))
        if keepDocx:
            print(f"POP DOCX 已生成: {docxPath}", flush=True)
            return str(docxPath)
        pdfPath = self.toPdf(docxPath)
        try:
            docxPath.unlink()
        except OSError:
            pass
        print(f"POP 已生成: {pdfPath}", flush=True)
        return str(pdfPath)

    def addBorder(self, parent, tag, val="dotted", sz="4", color="000000"):
        """向表格边框父节点追加单边样式"""
        el = OxmlElement(f"w:{tag}")
        el.set(qn("w:val"), val)
        el.set(qn("w:color"), color)
        el.set(qn("w:sz"), sz)
        el.set(qn("w:space"), "0")
        parent.append(el)

    def setDashBorder(self, tblPr):
        """汇总表设置为虚线边框"""
        old = tblPr.find(qn("w:tblBorders"))
        if old is not None:
            tblPr.remove(old)
        borders = OxmlElement("w:tblBorders")
        for side in ("top", "left", "bottom", "right", "insideH", "insideV"):
            self.addBorder(borders, side)
        tblPr.append(borders)

    def setSpacing(self, pElem):
        """段前/段后归零，单倍行距，避免固定大行高"""
        pPr = pElem.find(qn("w:pPr"))
        if pPr is None:
            pPr = OxmlElement("w:pPr")
            pElem.insert(0, pPr)
        oldSp = pPr.find(qn("w:spacing"))
        if oldSp is not None:
            pPr.remove(oldSp)
        spacing = OxmlElement("w:spacing")
        spacing.set(qn("w:before"), "0")
        spacing.set(qn("w:after"), "0")
        spacing.set(qn("w:line"), "240")
        spacing.set(qn("w:lineRule"), "auto")
        pPr.append(spacing)

    def clearHeight(self, tr):
        """移除表格行固定高度"""
        trPr = tr.find(qn("w:trPr"))
        if trPr is None:
            return
        hr = trPr.find(qn("w:trHeight"))
        if hr is not None:
            trPr.remove(hr)

    def setCellSingle(self, tc, text, tcTemplate):
        """按模板单元格样式写入单行文本"""
        # 取模板单元格第一个段落作为样式来源
        templateParas = tcTemplate.findall(qn("w:p"))
        srcPara = templateParas[0] if templateParas else None
        # 清空目标单元格内容，保留后续重新写入的样式结构
        for child in list(tc):
            tc.remove(child)
        # 复制模板单元格属性，保持边框、宽度等样式
        tcPr = tcTemplate.find(qn("w:tcPr"))
        if tcPr is not None:
            tc.append(deepcopy(tcPr))
        if srcPara is not None:
            # 复制模板段落并替换其中的第一个文本节点
            newPara = deepcopy(srcPara)
            texts = newPara.findall(".//" + qn("w:t"))
            if texts:
                texts[0].text = text
                for node in texts[1:]:
                    node.text = ""
            # 统一段落行距，避免模板固定行高影响展示
            self.setSpacing(newPara)
            tc.append(newPara)

    def rebuildSummary(self, doc):
        """多 SKU 模板汇总表拆为 5 行 2 列"""
        oldTable = doc.tables[0]
        oldRow = oldTable.rows[0]
        # 备份左右单元格样式，后续新行复用
        leftTpl = deepcopy(oldRow.cells[0]._tc)
        rightTpl = deepcopy(oldRow.cells[1]._tc)
        tbl = oldTable._tbl
        tblPr = tbl.find(qn("w:tblPr"))

        # 清空原汇总表数据行
        for tr in list(tbl.findall(qn("w:tr"))):
            tbl.remove(tr)

        if tblPr is not None:
            self.setDashBorder(tblPr)

        # 按固定 5 个汇总字段重建行
        templateTr = deepcopy(oldRow._tr)
        for label, value in zip(self.summaryLabels, self.summaryValues):
            tr = deepcopy(templateTr)
            self.clearHeight(tr)
            leftTc, rightTc = tr.findall(qn("w:tc"))
            self.setCellSingle(leftTc, label, leftTpl)
            self.setCellSingle(rightTc, value, rightTpl)
            tbl.append(tr)

    def relaxRows(self, doc):
        """商品表去除固定行高并统一段落行距"""
        itemTbl = doc.tables[1]._tbl
        for tr in itemTbl.findall(qn("w:tr")):
            self.clearHeight(tr)
            for tc in tr.findall(qn("w:tc")):
                for p in tc.findall(qn("w:p")):
                    self.setSpacing(p)

    def setCellWidth(self, tc, width):
        """设置单元格宽度"""
        tcPr = tc.find(qn("w:tcPr"))
        if tcPr is None:
            tcPr = OxmlElement("w:tcPr")
            tc.insert(0, tcPr)
        tcW = tcPr.find(qn("w:tcW"))
        if tcW is None:
            tcW = OxmlElement("w:tcW")
            tcPr.append(tcW)
        tcW.set(qn("w:w"), str(width))
        tcW.set(qn("w:type"), "dxa")

    def setFloat(self, tbl, tblpX, tblpY):
        """表格相对页面绝对定位"""
        tblPr = tbl.find(qn("w:tblPr"))
        if tblPr is None:
            tblPr = OxmlElement("w:tblPr")
            tbl.insert(0, tblPr)

        # 先清理旧定位，避免重复定位节点
        for tag in ("tblpPr", "tblOverlap"):
            old = tblPr.find(qn(f"w:{tag}"))
            if old is not None:
                tblPr.remove(old)

        # 写入新的页面绝对定位参数
        tblp = OxmlElement("w:tblpPr")
        tblp.set(qn("w:leftFromText"), "180")
        tblp.set(qn("w:rightFromText"), "180")
        tblp.set(qn("w:vertAnchor"), "page")
        tblp.set(qn("w:horzAnchor"), "page")
        tblp.set(qn("w:tblpX"), str(tblpX))
        tblp.set(qn("w:tblpY"), str(tblpY))
        tblPr.insert(1, tblp)

        overlap = OxmlElement("w:tblOverlap")
        overlap.set(qn("w:val"), "never")
        tblPr.insert(2, overlap)

        # 移除普通缩进，避免与绝对定位叠加
        tblInd = tblPr.find(qn("w:tblInd"))
        if tblInd is not None:
            tblPr.remove(tblInd)

    def alignSignature(self, doc):
        """姓名表与签名表同 Y 绝对定位，并拆成两行与右侧等高对齐"""
        sigTbl = doc.tables[2]._tbl
        nameTbl = doc.tables[3]._tbl
        sigTblPr = sigTbl.find(qn("w:tblPr"))
        sigTblp = sigTblPr.find(qn("w:tblpPr")) if sigTblPr is not None else None
        if sigTblp is None:
            return

        # 姓名表跟随签名表的 Y 坐标对齐
        tblpY = sigTblp.get(qn("w:tblpY"), "10620")
        nameRows = nameTbl.findall(qn("w:tr"))
        if len(nameRows) == 1:
            nameTc = nameRows[0].find(qn("w:tc"))
            paras = nameTc.findall(qn("w:p")) if nameTc is not None else []
            sigRows = sigTbl.findall(qn("w:tr"))

            # 复用签名表两行结构，拆出姓名标题行和姓名内容行
            labelTr = deepcopy(sigRows[0])
            bodyTr = deepcopy(sigRows[1])

            # 清空原姓名表单行结构
            for tr in list(nameTbl.findall(qn("w:tr"))):
                nameTbl.remove(tr)

            # 写入姓名标题行
            labelTc = labelTr.find(qn("w:tc"))
            for p in list(labelTc.findall(qn("w:p"))):
                labelTc.remove(p)
            if paras:
                labelTc.append(deepcopy(paras[0]))
            self.setCellWidth(labelTc, self.nameTableWidth)

            # 写入姓名内容行
            bodyTc = bodyTr.find(qn("w:tc"))
            for p in list(bodyTc.findall(qn("w:p"))):
                bodyTc.remove(p)
            if len(paras) > 1:
                bodyTc.append(deepcopy(paras[1]))
            else:
                bodyTc.append(OxmlElement("w:p"))
            self.setCellWidth(bodyTc, self.nameTableWidth)

            nameTbl.append(labelTr)
            nameTbl.append(bodyTr)

            # 同步表格网格宽度，避免 Word 自动压缩姓名表
            tblGrid = nameTbl.find(qn("w:tblGrid"))
            if tblGrid is not None:
                for gridCol in tblGrid.findall(qn("w:gridCol")):
                    gridCol.set(qn("w:w"), str(self.nameTableWidth))

        self.setFloat(nameTbl, self.nameTableLeftX, tblpY)

    def checkTemplate(self):
        """只读检查多 SKU 模板结构，避免误覆盖服务商模板.docx"""
        if not self.multiTemplatePath.is_file():
            raise FileNotFoundError(f"多 SKU 模板不存在: {self.multiTemplatePath}")

        # 只读取模板结构，不复制、不重建、不保存，避免破坏已确认样式
        doc = Document(str(self.multiTemplatePath))
        if len(doc.tables) < 4:
            raise ValueError(f"多 SKU 模板表格数量不足，当前为 {len(doc.tables)} 个")

        summaryTable = doc.tables[0]
        itemTable = doc.tables[1]
        signatureTable = doc.tables[2]
        nameTable = doc.tables[3]

        signaturePr = signatureTable._tbl.find(qn("w:tblPr"))
        signatureFloat = signaturePr.find(qn("w:tblpPr")) if signaturePr is not None else None
        signatureAttrs = None
        if signatureFloat is not None:
            signatureAttrs = {key.split("}")[1]: value for key, value in signatureFloat.attrib.items()}

        summaryOk = (
            len(summaryTable.rows) == 1
            and len(summaryTable.columns) == 2
            and len(summaryTable.rows[0].cells[1].paragraphs) >= 5
        )
        itemOk = len(itemTable.rows) >= self.multiItemPresetRows + 1 and len(itemTable.columns) == 3
        signatureOk = signatureFloat is not None
        nameOk = len(nameTable.rows) == 1 and len(nameTable.columns) == 1

        print(f"多 SKU 模板检查: {self.multiTemplatePath}", flush=True)
        print(f"表格数量: {len(doc.tables)}", flush=True)
        print(
            f"汇总表: {len(summaryTable.rows)} 行 {len(summaryTable.columns)} 列，"
            f"右侧段落 {len(summaryTable.rows[0].cells[1].paragraphs)} 个",
            flush=True,
        )
        print(f"商品表: {len(itemTable.rows)} 行 {len(itemTable.columns)} 列", flush=True)
        print(f"签名表浮动定位: {signatureAttrs}", flush=True)
        print(f"姓名表: {len(nameTable.rows)} 行 {len(nameTable.columns)} 列", flush=True)

        if not (summaryOk and itemOk and signatureOk and nameOk):
            raise ValueError("多 SKU 模板结构不符合当前确认样式，请检查服务商模板.docx")
        return True


if __name__ == "__main__":
    config = {
        "mode": "build",
        "detailInfo": {
            "shipmentId": "FBA19TEST01",
            "shopName": "Lydia deal-CA",
            "name": "FBA STA test export",
            "fulfillmentCenterId": "YEG2",
            "createTime": "2026-06-17",
            "source": "JUNMALL CN Hangzhou Zhejiang 311100",
            "items": [
                {"fnSku": "X004YE2RTB", "msku": "2pc Toilet Flapper Red-CA", "quantity": 250},
                {"fnSku": "X004YEDA37", "msku": "Toilet Flush Lever White-CA", "quantity": 150},
            ],
        },
    }

    svc = PopExport()
    if config["mode"] == "prepareTemplate":
        svc.checkTemplate()
    else:
        savePath = svc.build(config["detailInfo"])
        print(f"调试输出: {savePath}")
