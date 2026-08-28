"""简历文件读取、OCR和基础字段提取。"""

import re
from pathlib import Path

import numpy as np
from docx import Document
from PIL import Image
from pypdf import PdfReader


class ResumeExtractor:
    """从图片、PDF或Word简历中提取可供人事复核的基本信息。"""

    def __init__(self):
        self.ocrEngine = None
        self.allowedSuffixes = {".png", ".jpg", ".jpeg", ".webp", ".bmp", ".pdf", ".docx"}

    def extract(self, filePath):
        """读取文件文本并返回字段和原始识别文字。"""
        path = Path(filePath)
        suffix = path.suffix.lower()
        if suffix not in self.allowedSuffixes:
            raise ValueError("仅支持 PNG、JPG、WEBP、BMP、PDF 和 DOCX 简历")
        if suffix == ".docx":
            text = self.readDocx(path)
        elif suffix == ".pdf":
            text = self.readPdf(path)
        else:
            text = self.readImage(path)
        normalized = self.normalize(text)
        if len(normalized) < 10:
            raise ValueError("未能从简历中识别出足够文字，请上传更清晰的文件")
        return {**self.parse(normalized), "sourceText": normalized}

    def readDocx(self, path):
        """读取 Word 简历正文和表格文本。"""
        document = Document(path)
        lines = [paragraph.text for paragraph in document.paragraphs]
        for table in document.tables:
            for row in table.rows:
                lines.append(" ".join(cell.text for cell in row.cells))
        return "\n".join(lines)

    def readPdf(self, path):
        """优先读取 PDF 文本，扫描件则逐页 OCR。"""
        reader = PdfReader(path)
        text = "\n".join(page.extract_text() or "" for page in reader.pages)
        if len(self.normalize(text)) >= 80:
            return text
        try:
            import fitz
        except ImportError as exc:
            raise RuntimeError("扫描版 PDF 需要安装 PyMuPDF") from exc
        lines = []
        with fitz.open(path) as document:
            for page in document:
                pixmap = page.get_pixmap(matrix=fitz.Matrix(2, 2), alpha=False)
                image = Image.frombytes("RGB", [pixmap.width, pixmap.height], pixmap.samples)
                lines.append(self.runOcr(image))
        return "\n".join(lines)

    def readImage(self, path):
        """OCR识别简历图片。"""
        with Image.open(path) as image:
            return self.runOcr(image.convert("RGB"))

    def runOcr(self, image):
        """使用 RapidOCR 在本地识别图片文字。"""
        if self.ocrEngine is None:
            try:
                from rapidocr_onnxruntime import RapidOCR
            except ImportError as exc:
                raise RuntimeError("图片识别组件未安装，请执行 pip install -r requirements.txt") from exc
            self.ocrEngine = RapidOCR()
        result, _ = self.ocrEngine(np.asarray(image))
        return "\n".join(item[1] for item in result or [])

    def normalize(self, text):
        """清理识别文字中的空白和空行。"""
        lines = [re.sub(r"[ \t]+", " ", line).strip() for line in text.replace("\r", "\n").split("\n")]
        return "\n".join(line for line in lines if line)

    def findLabeled(self, text, labels, pattern=r"[^\n]{1,30}"):
        """按中文标签查找字段值。"""
        labelPattern = "|".join(re.escape(label) for label in labels)
        match = re.search(rf"(?:{labelPattern})\s*[：:]?\s*({pattern})", text, re.IGNORECASE)
        return match.group(1).strip(" ：:") if match else ""

    def parse(self, text):
        """从简历文字中推断公开基本信息，所有结果仍需人事复核。"""
        emailMatch = re.search(r"[A-Z0-9._%+-]+@[A-Z0-9.-]+\.[A-Z]{2,}", text, re.IGNORECASE)
        phoneMatch = re.search(r"(?<!\d)1[3-9]\d{9}(?!\d)", text)
        name = self.findLabeled(text, ["姓名", "Name"], r"[\u4e00-\u9fa5·A-Za-z ]{2,20}")
        if not name:
            for line in text.splitlines()[:8]:
                if re.fullmatch(r"[\u4e00-\u9fa5·]{2,4}", line):
                    name = line
                    break
        education = next((item for item in ["博士", "硕士", "本科", "大专", "高中"] if item in text), "")
        school = self.findLabeled(text, ["毕业院校", "学校", "院校"], r"[^\n]{2,40}")
        if not school:
            school = next((line for line in text.splitlines() if "大学" in line or "学院" in line), "")
        major = self.findLabeled(text, ["专业", "所学专业"], r"[^\n]{2,30}")
        city = self.findLabeled(text, ["意向城市", "所在城市", "现居城市", "期望城市", "地点"], r"[^\n]{2,20}")
        position = self.findLabeled(text, ["意向岗位", "应聘职位", "目标岗位", "期望职位", "求职意向"], r"[^\n]{2,40}")
        position = re.split(r"\s*(?:意向城市|期望城市|所在城市|现居城市)\s*[：:]", position, maxsplit=1)[0].strip()
        return {
            "name": name,
            "email": emailMatch.group(0) if emailMatch else "",
            "phone": phoneMatch.group(0) if phoneMatch else "",
            "education": education,
            "school": school,
            "major": major,
            "city": city,
            "position": position,
        }
